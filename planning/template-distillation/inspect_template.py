"""Read-only inventory for the retained UM syllabus template."""

from pathlib import Path
from zipfile import ZipFile

from docx import Document
from docx.oxml.ns import qn


REFERENCE = Path(__file__).resolve().parents[2] / "resources" / "UM Syllabus Template.docx"


def paragraph_record(index, paragraph, part="document"):
    fmt = paragraph.paragraph_format
    return {
        "part": part,
        "index": index,
        "style": paragraph.style.name if paragraph.style else None,
        "text": paragraph.text,
        "alignment": str(paragraph.alignment),
        "left_indent": str(fmt.left_indent),
        "first_line_indent": str(fmt.first_line_indent),
        "space_before": str(fmt.space_before),
        "space_after": str(fmt.space_after),
        "line_spacing": str(fmt.line_spacing),
    }


doc = Document(REFERENCE)

print("PARAGRAPHS")
for i, p in enumerate(doc.paragraphs):
    print(paragraph_record(i, p))

print("RUNS")
for i, p in enumerate(doc.paragraphs):
    print(f"PARAGRAPH {i} style={p.style.name if p.style else None}")
    for ri, run in enumerate(p.runs):
        rpr = run._r.rPr
        print({
            "run": ri,
            "text": run.text,
            "font": run.font.name,
            "size": str(run.font.size),
            "bold": run.bold,
            "italic": run.italic,
            "color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else None,
            "hidden": run.font.hidden,
            "highlight": str(run.font.highlight_color),
            "drawing": bool(run._r.xpath('.//w:drawing')),
            "has_rPr": rpr is not None,
        })

print("TABLES")
for ti, table in enumerate(doc.tables):
    print(f"TABLE {ti}: rows={len(table.rows)} cols={len(table.columns)} style={table.style.name if table.style else None}")
    for ri, row in enumerate(table.rows):
        print(f"  ROW {ri}: {[cell.text for cell in row.cells]}")

for si, section in enumerate(doc.sections):
    print(f"HEADER {si}")
    for i, p in enumerate(section.header.paragraphs):
        print(paragraph_record(i, p, f"header{si}"))
    print(f"FOOTER {si}")
    for i, p in enumerate(section.footer.paragraphs):
        print(paragraph_record(i, p, f"footer{si}"))

print("STYLES")
for style in doc.styles:
    if style.type != 1:
        continue
    font = style.font
    pf = style.paragraph_format
    print({
        "name": style.name,
        "base": style.base_style.name if style.base_style else None,
        "font": font.name,
        "size": str(font.size),
        "bold": font.bold,
        "italic": font.italic,
        "color": str(font.color.rgb) if font.color and font.color.rgb else None,
        "alignment": str(pf.alignment),
        "before": str(pf.space_before),
        "after": str(pf.space_after),
        "line": str(pf.line_spacing),
        "left": str(pf.left_indent),
        "first": str(pf.first_line_indent),
        "keep_next": pf.keep_with_next,
    })

with ZipFile(REFERENCE) as package:
    print("PACKAGE")
    for info in package.infolist():
        print(f"{info.filename}\t{info.file_size}\t{info.CRC}")
