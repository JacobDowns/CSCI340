"""Build the CSCI 340 syllabus from the retained UM syllabus template.

The retained template controls the document's visual system and institutional
content. Course content is inserted into documented editable slots while the
logo, styles, numbering, policies, resources, hyperlinks, and footer remain
source-derived.
"""

from __future__ import annotations

import hashlib
import shutil
import time
from copy import deepcopy
from pathlib import Path
from zipfile import ZIP_DEFLATED, ZipFile

from docx import Document
from docx.enum.text import WD_BREAK
from docx.oxml.ns import qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph


ROOT = Path(__file__).resolve().parents[1]
REFERENCE = ROOT / "resources" / "UM Syllabus Template.docx"
OUTPUT = Path(__file__).resolve().parent / "CSCI340-syllabus-UM-template-Autumn-2026-draft.docx"
EXPECTED_REFERENCE_SHA256 = "0F644C7DD84FC04C464ADFE4D7031E68A20B4FCCA9B0BF1A53ED6373205A5A64"


CATALOG_DESCRIPTION = (
    "Fundamentals of data modeling, the relational mode, normal forms, file "
    "organization, index structures and SQL. Major project involving the "
    "design and implementation of a relational database."
)

OUTCOMES = [
    "Translate an ambiguous problem description into a conceptual and relational data model.",
    "Write, test, and explain SQL queries involving joins, aggregation, subqueries, common table expressions, and window functions.",
    "Use relational algebra and dependencies to reason about query meaning and schema quality.",
    "Select and enforce keys, constraints, and transaction boundaries that protect data integrity.",
    "Diagnose normalization problems and justify an appropriate decomposition or deliberate denormalization.",
    "Predict how storage structures and indexes may affect a workload, then test that prediction with query plans and measurements.",
    "Explain transaction anomalies and evaluate concurrency behavior using isolation and serializability concepts.",
    "Compare row-oriented, columnar, database, file, and DataFrame representations for an analytical task.",
    "Construct and evaluate spatial queries, joins, and indexes while accounting for coordinate reference systems.",
    "Communicate a defensible data-system design using requirements, evidence, tradeoffs, and limitations.",
]

UNIT_OBJECTIVES = [
    "Unit 1 - Represent and query: model requirements, write and test SQL, and connect queries to relational algebra.",
    "Unit 2 - Correctness and performance: analyze dependencies, transactions, indexes, storage, and query plans.",
    "Unit 3 - Analytical and spatial systems: compare PostgreSQL, DuckDB, Parquet, Polars, and PostGIS for appropriate workloads.",
    "Unit 4 - Choose and defend: integrate evidence into a reproducible database design and architecture memo.",
]

SCHEDULE = [
    (
        "Week 1 (August 24-28) - Why databases?",
        "Relations, schemas, data independence, and the limits of file-based workflows. Studio: diagnose a messy scientific dataset. Group project: dataset and stakeholder sketch.",
    ),
    (
        "Week 2 (August 31-September 4) - SQL foundations",
        "Selection, projection, filtering, sorting, expressions, nulls, and query tests. Lab: turn natural-language questions into small SQL queries and boundary cases.",
    ),
    (
        "Week 3 (September 7-11) - Joins and aggregation",
        "Labor Day is Monday, September 7; no class. Joins, grouping, aggregation, and query decomposition. Capstone: candidate questions and initial workload.",
    ),
    (
        "Week 4 (September 14-18) - Relational algebra and equivalence",
        "Relational operators, compositional reasoning, equivalence, and logical plans. Studio: translate SQL to algebra and justify meaning-preserving transformations.",
    ),
    (
        "Week 5 (September 21-25) - ER modeling foundations",
        "Entities, attributes, relationships, identifiers, cardinality, participation, and diagram conventions. Studio: convert a narrative specification into competing ER models.",
    ),
    (
        "Week 6 (September 28-October 2) - ER modeling in depth",
        "Weak entities, associative entities, recursive relationships, specialization/generalization, temporal requirements, and ambiguous business rules. Studio: critique and revise ER diagrams.",
    ),
    (
        "Week 7 (October 5-9) - From ER models to relational schemas",
        "Map entities and relationships to relations; select keys; express domain, entity, and referential constraints. Capstone: requirements, ER diagram, and relational mapping review.",
    ),
    (
        "Week 8 (October 12-16) - Dependencies and normalization",
        "Indigenous People's Day is Monday, October 12; no class. Functional dependencies, closure, candidate keys, anomalies, 3NF, BCNF, lossless decomposition, and dependency preservation.",
    ),
    (
        "Week 9 (October 19-23) - Review and midterm exam",
        "Synthesize SQL, relational algebra, ER modeling, relational mapping, constraints, and normalization. Midterm exam: Friday, October 23, during the regular class meeting.",
    ),
    (
        "Week 10 (October 26-30) - Advanced SQL",
        "Common table expressions, windows, conditional aggregation, views, and recursive queries as time permits. Capstone: PostgreSQL schema, ingestion, constraints, and representative queries.",
    ),
    (
        "Week 11 (November 2-6) - Storage, indexes, and query plans",
        "Pages, records, B-trees, hashing, composite indexes, selectivity, scans, joins, cardinality estimates, and plan choices. Experiment: predict, measure, and explain index effects.",
    ),
    (
        "Week 12 (November 9-13) - Transactions and concurrency",
        "Veterans Day is Wednesday, November 11; no class. ACID, schedules, anomalies, serializability, isolation, locking, and MVCC. Capstone: correctness and transaction memo.",
    ),
    (
        "Week 13 (November 16-20) - Analytical databases and DataFrames",
        "A condensed comparison of OLTP and OLAP, row and column storage, Parquet, DuckDB, data movement, and Polars lazy expressions. Lab: one analytical question across SQL and DataFrame interfaces.",
    ),
    (
        "Week 14 (November 23-27) - Spatial database foundations",
        "Wednesday is a non-instructional travel day; Thursday and Friday are the Thanksgiving holiday. The shortened week introduces spatial types, coordinate reference systems, and spatial predicates.",
    ),
    (
        "Week 15 (November 30-December 4) - Spatial joins, integration, and review",
        "PostGIS spatial joins and indexes, selected broader database architectures, group-project design defense, and cumulative final review. Friday, December 4 is the last day of regular classes.",
    ),
    (
        "Final Exams (December 7-11)",
        "Provisional final exam date: Wednesday, December 9. The exact time must match the registrar's final-exam schedule once the regular class meeting time is known. Autumn Commencement is Friday, December 11.",
    ),
]


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest().upper()


def open_document_with_retry(path: Path, attempts: int = 8) -> Document:
    """Tolerate the short Windows file-indexing lock that can follow a copy."""
    for attempt in range(attempts):
        try:
            return Document(path)
        except PermissionError:
            if attempt == attempts - 1:
                raise
            time.sleep(0.25)
    raise AssertionError("unreachable")


def clear_paragraph(paragraph: Paragraph) -> None:
    for child in list(paragraph._p):
        if child.tag != qn("w:pPr"):
            paragraph._p.remove(child)


def set_paragraph(
    paragraph: Paragraph,
    text: str,
    *,
    style: str | None = None,
    bold: bool | None = None,
    italic: bool | None = None,
    size: float | None = None,
) -> Paragraph:
    clear_paragraph(paragraph)
    if style is not None:
        paragraph.style = style
    run = paragraph.add_run(text)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    return paragraph


def set_labeled_paragraph(paragraph: Paragraph, label: str, value: str, *, size=12) -> Paragraph:
    clear_paragraph(paragraph)
    label_run = paragraph.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(size)
    value_run = paragraph.add_run(value)
    value_run.font.size = Pt(size)
    return paragraph


def remove_paragraph(paragraph: Paragraph) -> None:
    parent = paragraph._element.getparent()
    parent.remove(paragraph._element)
    paragraph._p = paragraph._element = None


def clone_after(anchor: Paragraph, template: Paragraph, text: str, *, style=None, bold=None) -> Paragraph:
    element = deepcopy(template._p)
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    set_paragraph(paragraph, text, style=style, bold=bold)
    return paragraph


def insert_text_after(anchor: Paragraph, text: str, *, style="Normal", bold=None) -> Paragraph:
    element = deepcopy(anchor._p)
    anchor._p.addnext(element)
    paragraph = Paragraph(element, anchor._parent)
    set_paragraph(paragraph, text, style=style, bold=bold)
    return paragraph


def find_exact(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text == text:
            return paragraph
    raise KeyError(f"Paragraph not found: {text!r}")


def find_startswith(doc: Document, text: str) -> Paragraph:
    for paragraph in doc.paragraphs:
        if paragraph.text.startswith(text):
            return paragraph
    raise KeyError(f"Paragraph prefix not found: {text!r}")


def capture_slots(doc: Document) -> dict[str, object]:
    paragraphs = doc.paragraphs
    return {
        "logo_note": paragraphs[1],
        "title": find_exact(doc, "[Course Name] Syllabus"),
        "subtitle": find_exact(doc, "[Term Offered]"),
        "catalog_note": find_startswith(doc, "[Instructor Note: Consult"),
        "course_details_label": find_exact(doc, "Course Details"),
        "course_details": paragraphs[7:14],
        "instructor_label": find_startswith(doc, "Instructor Information"),
        "instructor_details": paragraphs[15:20],
        "course_info_blank": paragraphs[20],
        "description": find_startswith(doc, "Course Description:"),
        "outcomes_intro": find_startswith(doc, "Course Learning Outcomes:"),
        "outcomes": paragraphs[24:27],
        "objectives_intro": find_startswith(doc, "Module/Unit Learning Objectives:"),
        "objectives": paragraphs[28:32],
        "objectives_blank": paragraphs[32],
        "calendar_heading": find_startswith(doc, "Course Calendar"),
        "calendar_note": find_startswith(doc, "[Instructor Note: Provide a chronological"),
        "calendar_blank": paragraphs[35],
        "grading_intro": find_startswith(doc, "Grading Criteria"),
        "grading_blank": paragraphs[38],
        "grading_scale": find_startswith(doc, "Grading Scale"),
        "attendance": find_startswith(doc, "[Instructor Note: Add a description of your attendance"),
        "ai": find_startswith(doc, "[Instructor Note: All instructors should include"),
        "workload": find_startswith(doc, "[Instructor Note: Include a description of the expected amount"),
        "communication": find_startswith(doc, "[Instructor Note: Include a description of netiquette"),
        "academic_heading": find_exact(doc, "Academic Misconduct and the Student Conduct Code"),
        "ode_blank": paragraphs[54],
        "health": find_startswith(doc, "[Instructor Note: Add a description of any safety"),
        "inclusion_note": find_startswith(doc, "[Instructor Note: Include a school-specific"),
        "heading3_template": find_exact(doc, "Attendance Policy"),
        "heading4_template": find_exact(doc, "Food Pantry Program"),
        "normal_template": find_exact(doc, "The University of Montana acknowledges that we are in the aboriginal territories of the Salish and Kalispel people. Today, we honor the path they have always shown us in caring for this place for the generations to come."),
    }


def fill_course_information(doc: Document, slots: dict[str, object]) -> None:
    remove_paragraph(slots["logo_note"])
    set_paragraph(slots["title"], "Database Design Syllabus")
    set_paragraph(slots["subtitle"], "Autumn 2026 - Development Draft")
    remove_paragraph(slots["catalog_note"])
    set_paragraph(slots["course_details_label"], "Course Details", bold=True)

    details = [
        ("Course Title", "Database Design"),
        ("Course Number", "CSCI 340"),
        ("Course Credits", "3"),
        ("Course Offering", "Fall"),
        ("Prerequisite", "CSCI 232"),
        ("Course Meeting Days/Times", "To be scheduled"),
        ("Course Location", "To be scheduled"),
        ("Course Format", "Published with the final course offering"),
        ("Course Textbook and/or Resources", "No single required textbook in this development draft; assigned readings combine course notes, selected references, and official PostgreSQL, PostGIS, DuckDB, Polars, Parquet, and Quarto documentation."),
    ]
    source_detail = slots["course_details"][0]
    for paragraph, (label, value) in zip(slots["course_details"], details[:7]):
        set_labeled_paragraph(paragraph, label, value)
    anchor = slots["course_details"][-1]
    for label, value in details[7:]:
        anchor = clone_after(anchor, source_detail, "")
        set_labeled_paragraph(anchor, label, value)

    set_paragraph(slots["instructor_label"], "Instructor Information", bold=True)
    instructor = [
        ("Instructor Name", "To be announced with the final course offering"),
        ("Email", "Published with the final course offering"),
        ("Office Location", "Published with the final course offering"),
        ("Office Hours", "At least one scheduled hour per week; final time, location, and virtual option to be announced"),
        ("Teaching Assistant(s)", "To be determined"),
    ]
    for paragraph, (label, value) in zip(slots["instructor_details"], instructor):
        set_labeled_paragraph(paragraph, label, value)
    remove_paragraph(slots["course_info_blank"])


def fill_description_and_outcomes(doc: Document, slots: dict[str, object]) -> None:
    paragraph = slots["description"]
    clear_paragraph(paragraph)
    run = paragraph.add_run("Official catalog description: ")
    run.bold = True
    paragraph.add_run(CATALOG_DESCRIPTION)

    anchor = insert_text_after(
        paragraph,
        "Course emphasis: This offering retains the catalog foundations while extending them through evidence-based problem solving, transactions and query plans, analytical systems, DataFrames, and spatial databases. Examples rotate among scientific, transactional, and organizational settings.",
    )
    anchor = insert_text_after(anchor, "Course Approach and Technology", style="Heading 3")
    for item in [
        "PostgreSQL for relational design, constraints, transactions, indexes, and query plans.",
        "PostGIS for spatial data types, relationships, joins, coordinate systems, and spatial indexes.",
        "DuckDB and Parquet for local analytical SQL and columnar data.",
        "Polars and Quarto for lazy DataFrame expressions and reproducible computational narratives.",
    ]:
        anchor = clone_after(anchor, slots["outcomes"][0], item)

    set_paragraph(
        slots["outcomes_intro"],
        "Course Learning Outcomes: By the end of the course, students should be able to:",
        bold=True,
    )
    outcome_paragraphs = list(slots["outcomes"])
    for paragraph in outcome_paragraphs[1:]:
        remove_paragraph(paragraph)
    anchor = set_paragraph(outcome_paragraphs[0], OUTCOMES[0])
    for outcome in OUTCOMES[1:]:
        anchor = clone_after(anchor, outcome_paragraphs[0], outcome)

    set_paragraph(
        slots["objectives_intro"],
        "Module/Unit Learning Objectives: The course is organized around four cumulative units:",
        bold=True,
    )
    for paragraph, objective in zip(slots["objectives"], UNIT_OBJECTIVES):
        set_paragraph(paragraph, objective)
    remove_paragraph(slots["objectives_blank"])


def fill_calendar(doc: Document, slots: dict[str, object]) -> None:
    set_paragraph(slots["calendar_heading"], "Course Calendar", style="Heading 2")
    set_paragraph(
        slots["calendar_note"],
        "Student Orientation is August 17-21. Classes begin Monday, August 24; regular classes end Friday, December 4; and final examinations run December 7-11. Assignment due dates will be added before publication. The instructional sequence is:",
    )
    anchor = slots["calendar_blank"]
    for index, (week, details) in enumerate(SCHEDULE):
        if index == 0:
            anchor = set_paragraph(anchor, week, style="Heading 3")
        else:
            anchor = clone_after(anchor, slots["heading3_template"], week, style="Heading 3")
        anchor = clone_after(anchor, slots["normal_template"], details, style="Normal")


def fill_grading(doc: Document, slots: dict[str, object]) -> None:
    set_paragraph(
        slots["grading_intro"],
        "Grading Criteria: Grades reflect correctness, reasoning, evidence, communication, and revision. Detailed rubrics will accompany major assignments.",
    )
    anchor = slots["grading_blank"]
    grading_items = [
        "Participation - 10%: quizzes, brief in-class exercises, design critiques, predictions, and constructive contribution to studio work.",
        "Assignments - 50%: individual assignments and labs are 30%; the milestone-based group database project is 20%.",
        "Midterm and final examinations - 40%: the midterm is 20% and the cumulative final examination is 20%.",
    ]
    for index, item in enumerate(grading_items):
        if index == 0:
            anchor = set_paragraph(anchor, item, style="List Paragraph")
            # Clone the real outcome-list numbering pattern onto the reused blank.
            anchor._p.remove(anchor._p.pPr)
            anchor._p.insert(0, deepcopy(slots["outcomes"][0]._p.pPr))
        else:
            anchor = clone_after(anchor, slots["outcomes"][0], item)

    anchor = clone_after(anchor, slots["heading3_template"], "Group Database Project (20%)", style="Heading 3")
    anchor = clone_after(
        anchor,
        slots["normal_template"],
        "Students will work in small groups, normally two or three, to design and implement a relational database. The project is part of the Assignments category and will develop through a few connected milestones.",
        style="Normal",
    )
    for item in [
        "Phase 1 - Concept and conceptual design: purpose, users, requirements, business rules, and ER model.",
        "Phase 2 - Logical design and SQL schema: relations, keys, constraints, normalization decisions, and PostgreSQL DDL.",
        "Phase 3 - Implementation and data: a working database, suitable data, and reproducible setup and testing materials.",
        "Phase 4 - Queries, evaluation, and presentation: intended questions, important design choices, and a final demonstration and defense.",
    ]:
        anchor = clone_after(anchor, slots["outcomes"][0], item)

    anchor = clone_after(
        anchor,
        slots["normal_template"],
        "Milestone instructions will identify both shared and individual responsibilities. Individual preparation, reflection, contribution evidence, and the ability to explain the group's design may contribute to each student's project grade. Detailed deliverables, milestone dates, team-formation procedures, and rubrics will be provided separately.",
        style="Normal",
    )

    set_paragraph(slots["grading_scale"], "Grading Scale", bold=True)


def fill_policies(doc: Document, slots: dict[str, object]) -> None:
    set_paragraph(
        slots["attendance"],
        "Class meetings combine short explanations, worked examples, prediction tasks, debugging, studio work, and discussion. Students should attend and participate consistently. The final course offering will state procedures for excused absences and making up work that depends on in-class collaboration.",
    )
    set_paragraph(
        slots["ai"],
        "Unless an assignment prohibits their use, generative AI and coding assistants may be used when students disclose material use, verify output against requirements and authoritative documentation, preserve enough process evidence to explain the work, and remain able to defend every submitted design and result. AI use on the individual midterm is not permitted unless explicitly authorized. Fabricated citations, concealed use, or submission of unverified output is not acceptable.",
    )
    set_paragraph(
        slots["workload"],
        "In addition to scheduled class meetings, students should plan for approximately six hours of work per week for reading, problem solving, labs, project milestones, testing, and revision. Some project weeks may require redistributing that time across the semester.",
    )
    set_paragraph(
        slots["communication"],
        "Use the instructor's published email for private or individual matters and the course discussion space for questions that may help the class. Technical questions should include the question being answered, a minimal reproducible example, the observed and expected results, and what has already been tested. Never post passwords, connection strings, or restricted data. Response-time expectations will be published with the final offering.",
    )
    set_paragraph(
        slots["health"],
        "This is a computing course and requires no specialized protective equipment. Take reasonable screen and movement breaks, use ergonomically appropriate work habits, and follow posted laboratory rules. Use only course-approved systems, data, and credentials. Report exposed credentials or unintended access immediately rather than attempting further exploration.",
    )
    remove_paragraph(slots["ode_blank"])
    remove_paragraph(slots["inclusion_note"])

    anchor = slots["communication"]
    additions = [
        (
            "Collaboration and Attribution",
            "Discussion and peer feedback are encouraged unless an assignment says otherwise. Submitted work must identify collaborators and distinguish shared discussion from individually produced analysis. Cite reused code, prose, data, and outside sources.",
        ),
        (
            "Late Work and Revision",
            "The proposed policy provides a 48-hour grace window for routine assignments, followed by a predictable late penalty, with advance arrangements for significant circumstances. Capstone checkpoints may be revised after feedback when the revision is substantive and submitted by the stated deadline. Exact procedures will appear in the final offering.",
        ),
        (
            "Data Ethics, Privacy, and Security",
            "Use only course-approved data and credentials. Do not upload restricted, personally identifiable, or license-incompatible data to unapproved services. Document provenance, minimize collected data, and consider how definitions, missingness, linkage, geographic resolution, and access decisions may cause harm.",
        ),
        (
            "Getting Help",
            "Use office hours, the course discussion space, and peer studio time early. Asking a precise question, documenting an unsuccessful approach, and revising after evidence are expected parts of database problem solving.",
        ),
    ]
    for heading, body in additions:
        anchor = clone_after(anchor, slots["heading3_template"], heading, style="Heading 3")
        anchor = clone_after(anchor, slots["normal_template"], body, style="Normal")


def audit_content(doc: Document) -> None:
    all_text = "\n".join(paragraph.text for paragraph in doc.paragraphs)
    assert "[Instructor Note:" not in all_text
    assert "[Course Name]" not in all_text
    assert "[Term Offered]" not in all_text
    assert "Database Design Syllabus" in all_text
    assert CATALOG_DESCRIPTION in all_text
    assert "Course Offering: Fall" in all_text
    assert "Prerequisite: CSCI 232" in all_text
    assert "individual assignments and labs are 30%" in all_text
    assert "Group Database Project (20%)" in all_text
    assert "Phase 4 - Queries, evaluation, and presentation" in all_text
    assert "Phase 5" not in all_text
    assert len([p for p in doc.paragraphs if p.style.name == "Heading 2"]) >= 7
    assert len(doc.inline_shapes) == 1
    section = doc.sections[0]
    assert section.page_width.inches == 8.5
    assert section.page_height.inches == 11
    assert section.left_margin.inches == 1
    assert section.right_margin.inches == 1
    assert "Doc ID: LEAD 08.15.2025 (v1)" in "\n".join(
        p.text for p in section.footer.paragraphs
    )


def package_hashes(path: Path) -> dict[str, str]:
    result = {}
    with ZipFile(path) as package:
        for name in package.namelist():
            result[name] = hashlib.sha256(package.read(name)).hexdigest()
    return result


def restore_reference_package_parts(reference: Path, edited: Path) -> None:
    """Keep only intended XML edits and restore every other source part byte-for-byte."""
    editable = {"word/document.xml", "docProps/core.xml"}
    temporary = edited.with_name(f"{edited.stem}.repacked.tmp{edited.suffix}")
    with ZipFile(reference) as source, ZipFile(edited) as generated:
        generated_parts = {name: generated.read(name) for name in editable}
        with ZipFile(temporary, "w", compression=ZIP_DEFLATED) as destination:
            for info in source.infolist():
                data = generated_parts.get(info.filename, source.read(info.filename))
                destination.writestr(info, data)
    temporary.replace(edited)


def verify_preservation(reference: Path, final: Path) -> list[str]:
    before = package_hashes(reference)
    after = package_hashes(final)
    preserve = [name for name in before if name not in {"word/document.xml", "docProps/core.xml"}]
    changed = [name for name in preserve if before.get(name) != after.get(name)]
    missing = [name for name in preserve if name not in after]
    if changed or missing:
        raise AssertionError(f"Preserve-only package mismatch: changed={changed}, missing={missing}")
    return sorted(set(before) - set(after))


def main() -> None:
    if sha256(REFERENCE) != EXPECTED_REFERENCE_SHA256:
        raise RuntimeError("The retained UM template changed; rerun template distillation before building.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    shutil.copy2(REFERENCE, OUTPUT)
    doc = open_document_with_retry(OUTPUT)
    slots = capture_slots(doc)
    fill_course_information(doc, slots)
    fill_description_and_outcomes(doc, slots)
    fill_calendar(doc, slots)
    fill_grading(doc, slots)
    fill_policies(doc, slots)

    core = doc.core_properties
    core.title = "CSCI 340 - Database Design Syllabus"
    core.subject = "University of Montana course syllabus development draft"
    core.author = "CSCI 340 course development project"
    core.keywords = "database design, SQL, PostgreSQL, PostGIS, DuckDB, Polars"
    core.comments = "Generated from the retained UM syllabus template and syllabus.qmd course design."
    audit_content(doc)
    doc.save(OUTPUT)
    restore_reference_package_parts(REFERENCE, OUTPUT)

    dropped = verify_preservation(REFERENCE, OUTPUT)
    print(OUTPUT)
    if dropped:
        print("Unreferenced package parts dropped by python-docx:", ", ".join(dropped))


if __name__ == "__main__":
    main()
